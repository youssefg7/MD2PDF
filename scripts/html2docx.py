import argparse
import os
import sys

import docx.styles
import docx.styles.style
import docx.text
import docx.text.paragraph
import docx.text.parfmt
from docx.shared import Inches

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from helpers import get_settings
from helpers.docx_utils import *
from helpers.utils import *
import docx

from bs4 import BeautifulSoup
import re

app_settings = get_settings()


if __name__ == "__main__":
    # parser = argparse.ArgumentParser(description="Convert html to docx")
    # parser.add_argument("input_html_path", type=str, help="Input html file")
    # args = parser.parse_args()
    # input_html_path = args.input_html_path


    input_html_path = "input-samples/Sample arabic.html"
    input_html_path = os.path.abspath(input_html_path)
    if not os.path.exists(app_settings.OUTPUT_DIR):
        os.makedirs(app_settings.OUTPUT_DIR)
    output_docx_file = (input_html_path.split("/")[-1]).split(".")[0] + ".docx"
    output_docx_path = os.path.join(app_settings.OUTPUT_DIR, output_docx_file)
    print(f"Converting {input_html_path} to {output_docx_path}...")

    with open(input_html_path, "r") as f:
        html = f.read()
    html = process_html(html)


    css_styles = read_css("assets/styles.css")
    html_template = """
<html>
<head>{css_styles}</head>
<body>{html_content}</body>
</html>
"""
    formatted_template = html_template.format(html_content=html, css_styles=css_styles)

    
    docx_title = (input_html_path.split("/")[-1]).split(".")[0]
    
    # docx generated from html html2docx package
    # buf = html2docx(html, title=docx_title)
    # with open(output_docx_path, "wb") as f:
    #     f.write(buf.getvalue())
    soup = BeautifulSoup(html, "html.parser")

    doc = docx.Document()
    styles = doc.styles
    styles : dict[docx.styles.style.BaseStyle]
    # Page break before heading 1
    h1 = styles["Heading 1"]
    h1_fmt = h1.paragraph_format
    h1_fmt.page_break_before = True
    
    first_h1 = True  # Track if it's the first h1 to avoid page break on the first one
    
    for tag in soup.find_all():
        if re.match(r"h[1-3]", tag.name):
            level = int(tag.name[-1])
            
            # Add page break before each <h1> except the first one
            if tag.name == "h1" and not first_h1:
                doc.add_page_break()
            if tag.name == "h1":
                first_h1 = False

            heading = doc.add_heading(tag.text, level=level)
            set_rtl_for_paragraph(heading)
        
        elif tag.name == "p":
            # Add paragraphs
            par = doc.add_paragraph()
            for content in tag.contents:
                if hasattr(content, 'name'):  # Handle inline elements like <strong>, <em>, <a>
                    handle_inline_styles(par, content)
                else:
                    par.add_run(content)
            set_rtl_for_paragraph(par)
        
        elif tag.name == "img" and tag.has_attr("src"):
            # Add images
            try:
                doc.add_picture(tag["src"], width=Inches(4))
            except Exception as e:
                print(f"Error adding image: {e}")
        
        elif tag.name == "ul" or tag.name == "ol":
            # Handle lists
            handle_list(doc, tag)
        
        # elif tag.name == "table":
        #     # Handle tables
        #     handle_table(doc, tag)
        
        elif tag.name == "br":
            # Handle line breaks
            doc.add_paragraph()  # Blank paragraph to represent line break

        elif tag.name == "div":
            # Add a division of text as a paragraph
            par = doc.add_paragraph(tag.text)
            set_rtl_for_paragraph(par)
            
        elif tag.name == "a":
            # Add a hyperlink
            par = doc.add_paragraph()  # Add a new paragraph for hyperlink
            add_hyperlink(par, tag["href"], tag.text)
            set_rtl_for_paragraph(par)
        
        elif tag.name == "blockquote":
            par = doc.add_paragraph(tag.text)
            par.paragraph_format.left_indent = Inches(0.5)  # Indent the blockquote
            set_rtl_for_paragraph(par)
    
    set_font_for_all_paragraphs(doc, "Arial")
    
    doc.save(output_docx_path)
    print(f"Docx saved to {output_docx_path}")
