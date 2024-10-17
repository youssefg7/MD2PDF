import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def set_rtl_for_paragraph(paragraph):
    """Set RTL direction for a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def set_rtl_for_table(table):
    """Set RTL direction for all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_rtl_for_paragraph(paragraph)


def handle_list(doc, list_tag):
    """Handle ordered (<ol>) and unordered (<ul>) lists."""
    if list_tag.name == "ul":
        for li in list_tag.find_all("li"):
            par = doc.add_paragraph(li.text, style="List Bullet")
            set_rtl_for_paragraph(par)
    elif list_tag.name == "ol":
        for li in list_tag.find_all("li"):
            par = doc.add_paragraph(li.text, style="List Number")
            set_rtl_for_paragraph(par)


def handle_table(doc, table_tag):
    """Handle table and its content safely."""
    rows = table_tag.find_all("tr")

    if not rows:
        print("Warning: No rows found in the table.")
        return

    # Get the number of columns from the first row
    first_row_cells = rows[0].find_all("td")
    num_cols = len(first_row_cells)

    # Create a table with the correct number of columns based on the first row
    table = doc.add_table(rows=1, cols=num_cols)

    # Add the first row to the table
    hdr_cells = table.rows[0].cells
    for i, td in enumerate(first_row_cells):
        hdr_cells[i].text = td.text

    # Iterate over remaining rows and add them to the table
    for row_idx, tr in enumerate(rows[1:], start=1):
        cols = tr.find_all("td")

        # Add a new row to the table
        row_cells = table.add_row().cells

        # Populate the new row with data
        for j in range(num_cols):
            if j < len(cols):
                row_cells[j].text = cols[j].text
            else:
                row_cells[j].text = (
                    ""  # Fill with empty string if there are fewer columns
                )

    # Apply RTL settings to the table
    set_rtl_for_table(table)


def handle_inline_styles(paragraph, tag):
    """Handle inline styles like <strong>, <em>, <a>, etc."""
    run = paragraph.add_run(tag.text)
    if tag.name == "strong" or tag.name == "b":
        run.bold = True
    # if tag.name == "em" or tag.name == "i":
    #     run.italic = True
    if tag.name == "a" and tag.has_attr("href"):
        run.text = f"{tag.text} ({tag['href']})"


def add_hyperlink(paragraph, url, text, color="0000EE", underline=True):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(
        docx.oxml.shared.qn("r:id"),
        r_id,
    )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement("w:r")

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement("w:color")
        c.set(docx.oxml.shared.qn("w:val"), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement("w:u")
        u.set(docx.oxml.shared.qn("w:val"), "none")
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def set_font_for_all_paragraphs(doc, font_name="Arial", font_size=12):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color (optional)
            # You can set italic, bold, etc., here as well
            # run.font.italic = True
            # run.font.bold = True
